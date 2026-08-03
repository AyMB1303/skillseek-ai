"""Extraction du texte d'un CV au format PDF.

Stratégie en cascade (S3-01) :
  1. Extraction directe de la couche texte du PDF — rapide et fidèle,
     suffisante pour la grande majorité des CV.
  2. Si le document ne contient pas de texte exploitable (CV scanné,
     c'est-à-dire une suite d'images), reconnaissance optique de
     caractères en solution de repli.

Les dépendances lourdes sont importées paresseusement : le module reste
utilisable (et testable) même si Tesseract n'est pas installé.
"""
import logging
import re

logger = logging.getLogger(__name__)

# En dessous de ce nombre de caracteres, on considere que la couche texte
# est absente ou inexploitable : le document est probablement scanne.
SEUIL_TEXTE_EXPLOITABLE = 120


class ResultatExtraction:
    """Texte extrait et informations sur la méthode employée."""

    def __init__(self, texte, methode, pages=0, erreur=None):
        self.texte = texte or ""
        self.methode = methode          # "texte_natif" | "ocr" | "echec"
        self.pages = pages
        self.erreur = erreur

    @property
    def reussie(self):
        return self.methode != "echec" and len(self.texte) >= SEUIL_TEXTE_EXPLOITABLE

    def to_dict(self):
        return {
            "methode": self.methode,
            "pages": self.pages,
            "caracteres": len(self.texte),
            "reussie": self.reussie,
            "erreur": self.erreur,
        }


def _extraire_couche_texte(chemin):
    """Extraction directe : rapide, sans perte, mais inopérante sur un scan."""
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber absent : extraction directe indisponible.")
        return "", 0

    morceaux = []
    with pdfplumber.open(chemin) as pdf:
        pages = len(pdf.pages)
        for page in pdf.pages:
            morceaux.append(page.extract_text() or "")
    return "\n".join(morceaux), pages


def _extraire_par_ocr(chemin):
    """Reconnaissance optique : chaque page est convertie en image puis lue."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        logger.warning("pytesseract ou pdf2image absent : OCR indisponible.")
        return "", 0

    images = convert_from_path(chemin, dpi=200)
    morceaux = [pytesseract.image_to_string(img, lang="fra+eng") for img in images]
    return "\n".join(morceaux), len(images)


def nettoyer(texte):
    """Normalise le texte brut avant analyse linguistique."""
    if not texte:
        return ""
    # Recolle les mots coupes par un tiret en fin de ligne
    texte = re.sub(r"-\s*\n\s*", "", texte)
    # Uniformise les espaces et les sauts de ligne multiples
    texte = re.sub(r"[ \t\xa0]+", " ", texte)
    texte = re.sub(r"\n{3,}", "\n\n", texte)
    # Retire les caracteres de controle residuels
    texte = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", texte)
    return texte.strip()


def _extraire_docx(chemin):
    """Documents Word : paragraphes et cellules de tableaux.

    Les tableaux sont fréquents dans les CV (mise en page en colonnes) et
    seraient perdus si l'on se limitait aux paragraphes.
    """
    try:
        import docx
    except ImportError:
        logger.warning("python-docx absent : lecture des .docx indisponible.")
        return "", 0

    document = docx.Document(chemin)
    morceaux = [p.text for p in document.paragraphs]
    for tableau in document.tables:
        for ligne in tableau.rows:
            morceaux.append(" | ".join(cellule.text.strip() for cellule in ligne.cells))
    return "\n".join(morceaux), 1


def extraire_texte(chemin):
    """Point d'entrée : renvoie un ResultatExtraction pour le CV donné."""
    # Documents Word : lecture directe, sans cascade
    if str(chemin).lower().endswith((".docx", ".doc")):
        try:
            texte, pages = _extraire_docx(chemin)
            texte = nettoyer(texte)
            if len(texte) >= SEUIL_TEXTE_EXPLOITABLE:
                return ResultatExtraction(texte, "docx", pages)
            return ResultatExtraction(
                texte, "echec", pages, "Le document Word ne contient pas de texte exploitable."
            )
        except Exception as exc:
            return ResultatExtraction(
                "", "echec", 0, f"Lecture du document Word impossible : {exc}"
            )

    # 1. Couche texte native
    try:
        texte, pages = _extraire_couche_texte(chemin)
        texte = nettoyer(texte)
        if len(texte) >= SEUIL_TEXTE_EXPLOITABLE:
            return ResultatExtraction(texte, "texte_natif", pages)
    except Exception as exc:  # fichier corrompu, protege par mot de passe...
        logger.info("Extraction directe impossible (%s), bascule vers l'OCR.", exc)
        pages = 0

    # 2. Repli sur la reconnaissance optique (CV scanne)
    try:
        texte_ocr, pages_ocr = _extraire_par_ocr(chemin)
        texte_ocr = nettoyer(texte_ocr)
        if len(texte_ocr) >= SEUIL_TEXTE_EXPLOITABLE:
            return ResultatExtraction(texte_ocr, "ocr", pages_ocr)
        return ResultatExtraction(
            texte_ocr,
            "echec",
            pages_ocr,
            "Le document ne contient pas de texte exploitable.",
        )
    except Exception as exc:
        logger.warning("OCR en échec : %s", exc)
        return ResultatExtraction("", "echec", pages, f"Lecture impossible : {exc}")
